"""文档解析层测试：txt/md/json/jsonl、anydoc 全格式解析与错误路径。

PDF/DOCX/CSV 测试文件在测试内动态生成（reportlab 写 PDF、python-docx 写段落表格），
不依赖仓库内的二进制 fixture。

PDF 用英文文本：anydoc 的 PDF 引擎对合成 PDF 的中文 CID 字体提取不可靠
（真实中文 PDF 的提取质量已由 学习文档/ 目录实测验证），单元测试只验证解析机制。
"""
from __future__ import annotations

import io

import pytest

from mcp.document_parser import (
    _ANYDOC_EXTENSIONS,
    SUPPORTED_EXTENSIONS,
    parse_document,
)

# ── 测试内生成 PDF/DOCX 的辅助 ─────────────────────────────────────────────

def make_pdf(pages_text: list[str]) -> bytes:
    """用 reportlab 生成含文本页的 PDF（Helvetica，ASCII 文本，可被 anydoc 提取）。"""
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    for text in pages_text:
        c.setFont("Helvetica", 12)
        c.drawString(72, 720, text)
        c.showPage()
    c.save()
    return buf.getvalue()


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """用 python-docx 生成含段落（可含表格）的 docx。"""
    import docx

    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        t = d.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                t.rows[i].cells[j].text = cell
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── txt / md ────────────────────────────────────────────────────────────────

def test_txt_whole_file_as_one_doc():
    docs = parse_document("新生指南.txt", "欢迎来到西电。".encode())
    assert len(docs) == 1
    assert docs[0]["title"] == "新生指南"
    assert docs[0]["content"] == "欢迎来到西电。"
    assert docs[0]["format"] == "txt"


def test_md_same_as_txt():
    docs = parse_document("排障手册.md", "# 标题\n正文内容".encode())
    assert docs[0]["title"] == "排障手册"
    assert docs[0]["format"] == "md"
    assert "# 标题\n正文内容" in docs[0]["content"]


def test_txt_empty_raises():
    with pytest.raises(ValueError, match="内容为空"):
        parse_document("空文档.txt", b"   \n ")


# ── json / jsonl ────────────────────────────────────────────────────────────

def test_json_array_docs():
    data = '[{"title": "校历", "content": "秋季学期开学。"}, {"title": "选课", "content": "分预选正选。"}]'.encode()
    docs = parse_document("知识.json", data)
    assert len(docs) == 2
    assert docs[0]["title"] == "校历"
    assert docs[0]["format"] == "json"          # 自动补 format
    assert docs[1]["content"] == "分预选正选。"


def test_json_fields_passthrough():
    data = '[{"title": "校历", "content": "内容", "domain": "affairs", "source_url": "https://x"}]'.encode()
    doc = parse_document("知识.json", data)[0]
    assert doc["domain"] == "affairs"
    assert doc["source_url"] == "https://x"


def test_json_not_array_raises():
    with pytest.raises(ValueError, match="数组格式"):
        parse_document("坏.json", b'{"title": "x", "content": "y"}')


def test_json_broken_raises():
    with pytest.raises(ValueError, match="JSON 解析失败"):
        parse_document("坏.json", b'[{"title": ')


def test_jsonl_line_docs():
    data = ('{"title": "校历", "content": "秋季学期开学。"}\n'
            '{"title": "选课", "content": "分预选正选。"}\n').encode()
    docs = parse_document("知识.jsonl", data)
    assert len(docs) == 2
    assert docs[0]["title"] == "校历"
    assert docs[0]["format"] == "jsonl"         # 自动补 format
    assert docs[1]["content"] == "分预选正选。"


def test_jsonl_skips_blank_lines():
    data = (b'{"title": "A", "content": "x"}\n'
            b'\n'
            b'{"title": "B", "content": "y"}\n')
    docs = parse_document("清单.jsonl", data)
    assert [d["title"] for d in docs] == ["A", "B"]


def test_jsonl_broken_line_raises():
    with pytest.raises(ValueError, match="JSONL 第 2 行解析失败"):
        parse_document("坏.jsonl", b'{"title": "A", "content": "x"}\n{"title": ')


def test_jsonl_non_object_line_raises():
    with pytest.raises(ValueError, match="JSONL 第 1 行应为对象"):
        parse_document("坏.jsonl", b'["not", "an", "object"]')


# ── pdf（anydoc）────────────────────────────────────────────────────────────

def test_pdf_multi_page_extracts_text():
    pdf = make_pdf(["Page one text. ", "Page two text. "])
    docs = parse_document("政策文件.pdf", pdf)
    assert len(docs) == 1
    doc = docs[0]
    assert doc["title"] == "政策文件"
    assert doc["format"] == "pdf"
    assert "Page one text" in doc["content"]
    assert "Page two text" in doc["content"]
    # 页码标注已随 pypdf 解析器移除（anydoc 输出 GFM Markdown，无 page_offsets）
    assert "page_offsets" not in doc


def test_pdf_scanned_no_text_layer_raises():
    pdf = make_pdf([" "])  # 空白页 ≈ 扫描件无文本层
    with pytest.raises(ValueError, match="扫描件|无文本层"):
        parse_document("扫描件.pdf", pdf)


def test_pdf_corrupt_raises():
    with pytest.raises(ValueError, match="PDF"):
        parse_document("损坏.pdf", b"%PDF-1.4\nthis is not a real pdf file")


# ── docx / csv（anydoc）─────────────────────────────────────────────────────

def test_docx_paragraphs_and_tables_in_order():
    docx = make_docx(
        ["第一条说明。", "第二条说明。"],
        table_rows=[["日期", "事项"], ["9月1日", "开学"]],
    )
    docs = parse_document("日程安排.docx", docx)
    assert len(docs) == 1
    text = docs[0]["content"]
    assert "第一条说明。" in text
    assert "第二条说明。" in text
    assert "| 日期 | 事项 |" in text            # anydoc 输出 GFM 表格，结构不丢
    assert "| 9月1日 | 开学 |" in text
    assert docs[0]["format"] == "docx"


def test_docx_empty_raises():
    empty = make_docx([])
    with pytest.raises(ValueError, match="空文档"):
        parse_document("空.docx", empty)


def test_csv_to_markdown_table():
    csv_bytes = "名称,价格\n奶茶,12\n".encode()
    docs = parse_document("价目表.csv", csv_bytes)
    assert len(docs) == 1
    text = docs[0]["content"]
    assert "| 名称 | 价格 |" in text
    assert "| 奶茶 | 12 |" in text
    assert docs[0]["format"] == "csv"


def test_legacy_office_extensions_accepted():
    # .doc/.xls/.ppt 等旧格式同样走 anydoc；这里验证扩展名→格式映射完整
    assert ".doc" in _ANYDOC_EXTENSIONS
    assert ".xls" in _ANYDOC_EXTENSIONS
    assert ".ppt" in _ANYDOC_EXTENSIONS


# ── 其他 ────────────────────────────────────────────────────────────────────

def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="不支持的文件格式"):
        parse_document("图片.png", b"whatever")


def test_no_extension_raises():
    with pytest.raises(ValueError, match="不支持的文件格式"):
        parse_document("README", b"text")


def test_supported_extensions():
    assert SUPPORTED_EXTENSIONS == {
        ".txt", ".md", ".json", ".jsonl",
    } | set(_ANYDOC_EXTENSIONS)
    # anydoc 声称覆盖的格式全部在支持列表内
    for ext in (".doc", ".docx", ".docm", ".ppt", ".pptx", ".xls", ".xlsx",
                ".xlsb", ".odt", ".ods", ".odp", ".rtf", ".epub", ".csv", ".pdf"):
        assert ext in SUPPORTED_EXTENSIONS
